#!/usr/bin/env python3
"""Build the publishable Enterprise AI Systems Delivery book."""

from __future__ import annotations

import re
import shutil
import subprocess
import hashlib
import os
import signal
import time
import zipfile
from pathlib import Path
from urllib.parse import quote, unquote

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from lxml import etree
from lxml import html as lxml_html
from pypdf import PdfReader


ROOT = Path(__file__).resolve().parents[1]
BOOK = ROOT
OUTPUT = BOOK / "output"
DOC_OUTPUT = OUTPUT / "doc"
PDF_OUTPUT = OUTPUT / "pdf"
TMP = ROOT / "tmp" / "publishing" / "enterprise-ai-system-delivery"
TITLE = "企业 AI 系统落地"
SUBTITLE = "从业务重构到可治理架构"
AUTHOR = "海波AI"
DATE = "2026 年 7 月"

PARTS = [
    (
        "第一篇　先判断什么值得做",
        "从工具诉求回到业务结果，确定项目边界、场景优先级、基线与价值假设。",
        [
            "manuscript/part-01-value/01-business-outcomes-not-ai-tools.md",
            "manuscript/part-01-value/02-system-boundaries.md",
            "manuscript/part-01-value/03-prioritization-baseline-roi.md",
        ],
    ),
    (
        "第二篇　重新设计人与工作的关系",
        "看清真实流程，重新安排人、AI 与系统的分工，并把异常、审批和责任写进目标流程。",
        [
            "manuscript/part-02-work-redesign/04-as-is-diagnosis.md",
            "manuscript/part-02-work-redesign/05-to-be-redesign.md",
            "manuscript/part-02-work-redesign/06-process-modeling.md",
            "manuscript/part-02-work-redesign/07-human-ai-system-roles.md",
        ],
    ),
    (
        "第三篇　把业务方案翻译成 AI 系统",
        "用四个视角连接知识、工作流、智能体与企业系统，形成可以共同讨论的解决方案。",
        [
            "manuscript/part-03-solution-architecture/08-four-plane-architecture.md",
            "manuscript/part-03-solution-architecture/09-choosing-software-rag-workflow-agent.md",
            "manuscript/part-03-solution-architecture/10-knowledge-context-engineering.md",
            "manuscript/part-03-solution-architecture/11-enterprise-integration-reliability.md",
            "manuscript/part-03-solution-architecture/12-architecture-review-package.md",
        ],
    ),
    (
        "第四篇　从解决方案到企业平台",
        "按任务和数据选择云端、本地或混合路径，并建立推理服务、模型网关与安全控制。",
        [
            "manuscript/part-04-platform-governance/13-cloud-local-hybrid.md",
            "manuscript/part-04-platform-governance/14-local-inference-engineering.md",
            "manuscript/part-04-platform-governance/15-model-gateway-control.md",
            "manuscript/part-04-platform-governance/16-data-identity-security-compliance.md",
        ],
    ),
    (
        "第五篇　从概念验证走向生产与长期运营",
        "把评估、追踪、服务目标、成本和组织责任组合成能够长期运行的生产机制。",
        [
            "manuscript/part-05-production-scale/17-evaluation-engineering.md",
            "manuscript/part-05-production-scale/18-llmops-ai-sre.md",
            "manuscript/part-05-production-scale/19-poc-to-production.md",
            "manuscript/part-05-production-scale/20-operating-model-90-day-roadmap.md",
        ],
    ),
]

CASES = [
    "cases/00-case-reading-guide.md",
    "cases/case-a-sales-knowledge-proposal-system.md",
    "cases/case-b-meeting-to-execution.md",
    "cases/case-c-pricing-contract-risk-assistant.md",
]

APPENDICES = [
    "appendices/A-project-toolkit.md",
    "appendices/B-glossary.md",
    "appendices/C-versioned-technical-cards.md",
    "appendices/D-code-companion.md",
    "appendices/E-references.md",
    "appendices/F-subject-index.md",
    "appendices/G-figure-list.md",
    "appendices/H-technical-deep-dives.md",
    "appendices/I-chapter-workbook.md",
]

FINAL_MARKDOWN_NAME = "企业AI系统落地-第四稿.md"
FINAL_DOCX_NAME = "企业AI系统落地-第四稿.docx"
FINAL_PDF_NAME = "企业AI系统落地-第四稿.pdf"


def run(command: list[str]) -> None:
    subprocess.run(command, check=True, cwd=ROOT)


def demote_headings(text: str) -> str:
    output: list[str] = []
    in_fence = False
    for line in text.splitlines():
        if line.startswith("```"):
            in_fence = not in_fence
        if not in_fence and re.match(r"^#{1,5} ", line):
            line = "#" + line
        output.append(line)
    return "\n".join(output).strip() + "\n"


def read(relative: str, *, demote: bool = False) -> str:
    text = (BOOK / relative).read_text(encoding="utf-8")
    # Chapter image paths are authored relative to manuscript/part-* but the
    # compiled Markdown is emitted under output/. Keep the published source
    # portable instead of baking absolute workstation paths into the book.
    if relative.startswith("manuscript/part-"):
        text = text.replace("](../../assets/", "](../assets/")
    return demote_headings(text) if demote else text.strip() + "\n"


def content_sources() -> list[str]:
    chapters = [chapter for _, _, items in PARTS for chapter in items]
    return [
        "manuscript/00-preface.md",
        "manuscript/00-prologue.md",
        *chapters,
        *CASES,
        *(appendix for appendix in APPENDICES if not appendix.endswith("G-figure-list.md")),
    ]


def generate_figure_list() -> Path:
    """Create a stable chapter-level list from all authored image references."""

    image_pattern = re.compile(r"!\[([^\]]+)\]\(([^)]+)\)")
    rows: list[tuple[str, str, str]] = []
    for relative in content_sources():
        text = (BOOK / relative).read_text(encoding="utf-8")
        heading = next(
            (line[2:].strip() for line in text.splitlines() if line.startswith("# ")),
            relative,
        )
        for caption, target in image_pattern.findall(text):
            rows.append((str(len(rows) + 1), caption.strip(), heading))

    lines = [
        "# 附录 G：插图目录",
        "",
        "本目录由出版构建脚本依据书稿图片引用生成。编号表示在定稿书稿中的出现顺序；章节内正式图号仍以图题为准。",
        "",
        "| 顺序 | 图题 | 位置 |",
        "|---:|---|---|",
    ]
    lines.extend(f"| {number} | {caption} | {heading} |" for number, caption, heading in rows)
    lines.extend(["", f"共 {len(rows)} 张插图。", ""])
    output = BOOK / "appendices" / "G-figure-list.md"
    output.write_text("\n".join(lines), encoding="utf-8")
    return output


def build_markdown() -> Path:
    OUTPUT.mkdir(parents=True, exist_ok=True)
    generate_figure_list()
    pieces = [
        "---\n",
        f'title: "{TITLE}"\n',
        f'subtitle: "{SUBTITLE}"\n',
        f'author: "{AUTHOR}"\n',
        f'date: "{DATE}"\n',
        'lang: "zh-CN"\n',
        "---\n\n",
        read("publication/00-publication-note.md"),
        "\n",
        read("publication/01-editorial-and-rights.md"),
        "\n",
        read("manuscript/00-preface.md"),
        "\n",
        read("manuscript/00-prologue.md"),
        "\n",
    ]
    for title, introduction, chapters in PARTS:
        pieces.extend([f"# {title}\n\n", f"{introduction}\n\n"])
        for chapter in chapters:
            pieces.extend([demote_headings(read(chapter)), "\n"])

    pieces.extend(
        [
            "# 案例篇\n\n",
            "三个案例分别展示完整系统、快速流程闭环与高风险控制。\n\n",
        ]
    )
    for case in CASES:
        pieces.extend([demote_headings(read(case)), "\n"])
    for appendix in APPENDICES:
        pieces.extend([read(appendix), "\n"])

    output = OUTPUT / FINAL_MARKDOWN_NAME
    output.write_text("".join(pieces), encoding="utf-8")
    return output


def first_heading(relative: str) -> str:
    for line in read(relative).splitlines():
        if line.startswith("# "):
            return line[2:].strip()
    raise ValueError(f"missing level-one heading: {relative}")


def build_docx_markdown(markdown: Path) -> Path:
    """Add a static contents page so DOCX and PDF work without refreshing fields."""
    lines = ["# 目录", "", "- 前言：企业 AI 的难点，从来不只在 AI", "- 序章：一场没有答案的企业 AI 方案会"]
    for title, _, chapters in PARTS:
        lines.append(f"- {title}")
        lines.extend(f"  - {first_heading(chapter)}" for chapter in chapters)
    lines.append("- 案例篇")
    lines.extend(f"  - {first_heading(case)}" for case in CASES)
    lines.append("- 附录")
    lines.extend(f"  - {first_heading(appendix)}" for appendix in APPENDICES)
    lines.append("")

    source = markdown.read_text(encoding="utf-8")
    metadata, body = source.split("---\n\n", 1)
    output = TMP / "book-docx.md"
    output.write_text(metadata + "---\n\n" + "\n".join(lines) + "\n" + body, encoding="utf-8")
    return output


def set_east_asia_font(style, font_name: str) -> None:
    style.font.name = font_name
    style._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def create_reference_docx(path: Path) -> None:
    default_reference = subprocess.run(
        ["pandoc", "--print-default-data-file", "reference.docx"],
        check=True,
        capture_output=True,
    ).stdout
    path.write_bytes(default_reference)
    document = Document(path)

    styles = document.styles
    normal = styles["Normal"]
    set_east_asia_font(normal, "Noto Sans SC")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(7)

    for style_name in ("Body Text", "First Paragraph"):
        if style_name in styles:
            style = styles[style_name]
            set_east_asia_font(style, "Noto Sans SC")
            style.font.size = Pt(10.5)
            style.paragraph_format.line_spacing = 1.5

    heading_sizes = {"Title": 30, "Subtitle": 16, "Heading 1": 24, "Heading 2": 18, "Heading 3": 14, "Heading 4": 11}
    for style_name, size in heading_sizes.items():
        if style_name not in styles:
            continue
        style = styles[style_name]
        set_east_asia_font(style, "Noto Sans SC")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor(15, 89, 97)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.space_before = Pt(16)
        style.paragraph_format.space_after = Pt(8)
        if style_name in ("Heading 1", "Heading 2"):
            style.paragraph_format.page_break_before = True

    for style_name in ("Title", "Subtitle", "Author", "Date"):
        if style_name in styles:
            styles[style_name].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for style_name in ("Caption", "Image Caption"):
        if style_name in styles:
            style = styles[style_name]
            set_east_asia_font(style, "Noto Sans SC")
            style.font.size = Pt(8.5)
            style.font.color.rgb = RGBColor(89, 100, 110)
            style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            style.paragraph_format.space_before = Pt(3)
            style.paragraph_format.space_after = Pt(8)
    if "Title" in styles:
        styles["Title"].paragraph_format.space_before = Pt(120)
    document.save(path)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, end])


def set_table_borders(table) -> None:
    """Apply borders without relying on a localized Word style name."""
    table_properties = table._tbl.tblPr
    borders = table_properties.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        table_properties.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "B7C7C9")


def polish_docx(path: Path) -> None:
    document = Document(path)
    settings = document.settings._element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")
    for section in document.sections:
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.1)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.0)
        section.header_distance = Cm(0.9)
        section.footer_distance = Cm(0.9)
        section.different_first_page_header_footer = True

        header = section.header.paragraphs[0]
        header.text = TITLE
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if header.runs:
            header.runs[0].font.size = Pt(8)
            header.runs[0].font.color.rgb = RGBColor(100, 110, 116)

        footer = section.footer.paragraphs[0]
        add_page_number(footer)

    for table in document.tables:
        set_table_borders(table)
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_after = Pt(2)
                    for run_item in paragraph.runs:
                        run_item.font.size = Pt(8.5)

    max_image_width = Cm(16.2)
    for shape in document.inline_shapes:
        if shape.width > max_image_width:
            scale = max_image_width / shape.width
            shape.height = int(shape.height * scale)
            shape.width = max_image_width

    for paragraph in document.paragraphs:
        if paragraph._p.xpath(".//w:drawing"):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_before = Pt(8)
            paragraph.paragraph_format.space_after = Pt(3)
        if paragraph.style.name in ("Caption", "Image Caption"):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.save(path)


def build_docx(markdown: Path) -> Path:
    DOC_OUTPUT.mkdir(parents=True, exist_ok=True)
    TMP.mkdir(parents=True, exist_ok=True)
    reference = TMP / "reference.docx"
    create_reference_docx(reference)
    docx_markdown = build_docx_markdown(markdown)
    output = DOC_OUTPUT / FINAL_DOCX_NAME
    run(
        [
            "pandoc",
            str(docx_markdown),
            "--from=gfm+yaml_metadata_block+implicit_figures+footnotes",
            f"--reference-doc={reference}",
            f"--resource-path={OUTPUT}",
            "--metadata=pagetitle:企业 AI 系统落地",
            "-o",
            str(output),
        ]
    )
    polish_docx(output)
    return output


def chrome_print(html: Path, pdf: Path) -> None:
    chrome = Path("/Applications/Google Chrome.app/Contents/MacOS/Google Chrome")
    pdf.unlink(missing_ok=True)
    profile = TMP / f"chrome-{pdf.stem}"
    if profile.exists():
        shutil.rmtree(profile)
    command = [
        str(chrome),
        "--headless=new",
        "--disable-gpu",
        "--disable-dev-shm-usage",
        "--disable-background-networking",
        "--disable-component-update",
        "--disable-default-apps",
        "--disable-extensions",
        "--disable-sync",
        "--metrics-recording-only",
        "--no-first-run",
        f"--user-data-dir={profile}",
        "--no-pdf-header-footer",
        f"--print-to-pdf={pdf}",
        html.resolve().as_uri(),
    ]
    # Chrome occasionally finishes writing a valid PDF but leaves its headless
    # parent process alive. Watch the output itself instead of waiting forever.
    process = subprocess.Popen(
        command,
        cwd=ROOT,
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        start_new_session=True,
    )
    deadline = time.monotonic() + 180
    last_size = -1
    stable_since: float | None = None
    exited_since: float | None = None
    complete = False
    try:
        while time.monotonic() < deadline:
            if pdf.exists():
                size = pdf.stat().st_size
                if size > 10_000 and size == last_size:
                    stable_since = stable_since or time.monotonic()
                    if time.monotonic() - stable_since >= 2:
                        try:
                            complete = len(PdfReader(str(pdf), strict=False).pages) > 0
                        except Exception:
                            complete = False
                        if complete:
                            break
                else:
                    last_size = size
                    stable_since = None
            if process.poll() is not None and not pdf.exists():
                exited_since = exited_since or time.monotonic()
                if time.monotonic() - exited_since >= 10:
                    break
            time.sleep(0.5)
    finally:
        if process.poll() is None:
            try:
                os.killpg(process.pid, signal.SIGTERM)
                process.wait(timeout=5)
            except (ProcessLookupError, subprocess.TimeoutExpired):
                try:
                    os.killpg(process.pid, signal.SIGKILL)
                except ProcessLookupError:
                    pass
                process.wait(timeout=5)
    if not complete:
        raise RuntimeError(f"Chrome did not finish a valid PDF at {pdf}")


def toc_page_map(html: Path, pdf: Path) -> dict[str, int]:
    tree = lxml_html.fromstring(html.read_text(encoding="utf-8"))
    reader = PdfReader(str(pdf))
    mapping: dict[str, int] = {}
    for link in tree.xpath('//nav[@id="TOC"]//a'):
        href = link.get("href")
        if not href or not href.startswith("#"):
            continue
        identifier = unquote(href[1:])
        destination_key = "/" + quote(identifier, safe="")
        destination = reader.named_destinations.get(destination_key)
        if destination is None:
            raise KeyError(f"missing PDF destination for TOC link {href}")
        mapping[identifier] = reader.get_destination_page_number(destination) + 1
    return mapping


def write_numbered_toc_html(
    source: Path,
    output: Path,
    mapping: dict[str, int],
) -> None:
    tree = lxml_html.fromstring(source.read_text(encoding="utf-8"))
    for link in tree.xpath('//nav[@id="TOC"]//a'):
        href = link.get("href")
        identifier = unquote(href[1:]) if href and href.startswith("#") else ""
        if identifier not in mapping:
            continue
        for old in link.xpath('./span[contains(concat(" ", @class, " "), " toc-page ")]'):
            link.remove(old)
        page = etree.SubElement(link, "span")
        page.set("class", "toc-page")
        page.text = str(mapping[identifier])
    output.write_text(
        etree.tostring(
            tree,
            encoding="unicode",
            method="html",
            doctype="<!DOCTYPE html>",
        ),
        encoding="utf-8",
    )


def build_pdf(markdown: Path) -> Path:
    PDF_OUTPUT.mkdir(parents=True, exist_ok=True)
    TMP.mkdir(parents=True, exist_ok=True)
    output = PDF_OUTPUT / FINAL_PDF_NAME
    source_html = TMP / "book.html"
    run(
        [
            "pandoc",
            str(markdown),
            "--from=gfm+yaml_metadata_block+implicit_figures+footnotes",
            "--standalone",
            "--section-divs",
            "--toc",
            "--toc-depth=2",
            f"--css={BOOK / 'publication' / 'book.css'}",
            f"--resource-path={OUTPUT}",
            "--embed-resources",
            "--metadata=toc-title:目录",
            "--metadata=pagetitle:企业 AI 系统落地",
            "-o",
            str(source_html),
        ]
    )

    current_pdf = TMP / "toc-pass-0.pdf"
    chrome_print(source_html, current_pdf)
    current_map = toc_page_map(source_html, current_pdf)

    candidate_pdf = current_pdf
    for iteration in range(1, 5):
        numbered_html = TMP / f"book-toc-pass-{iteration}.html"
        candidate_pdf = TMP / f"toc-pass-{iteration}.pdf"
        write_numbered_toc_html(source_html, numbered_html, current_map)
        chrome_print(numbered_html, candidate_pdf)
        candidate_map = toc_page_map(numbered_html, candidate_pdf)
        if candidate_map == current_map:
            shutil.copyfile(candidate_pdf, output)
            break
        current_map = candidate_map
    else:
        shutil.copyfile(candidate_pdf, output)
    # Chrome preserves the title but does not populate the PDF author field.
    run(
        [
            "exiftool",
            "-overwrite_original",
            f"-Author={AUTHOR}",
            f"-Title={TITLE}",
            str(output),
        ]
    )
    return output


def build_pdf_from_docx(docx: Path) -> Path:
    """Use LibreOffice as a layout-faithful fallback for long embedded HTML."""
    PDF_OUTPUT.mkdir(parents=True, exist_ok=True)
    output = PDF_OUTPUT / FINAL_PDF_NAME
    output.unlink(missing_ok=True)
    profile = TMP / "libreoffice-profile"
    command_prefix: list[str] = []
    fontconfig = Path("/opt/homebrew/etc/fonts/fonts.conf")
    if fontconfig.exists():
        command_prefix = ["env", f"FONTCONFIG_FILE={fontconfig}"]
    run(
        command_prefix
        + [
            "soffice",
            f"-env:UserInstallation={profile.resolve().as_uri()}",
            "--headless",
            "--convert-to",
            "pdf",
            "--outdir",
            str(PDF_OUTPUT),
            str(docx),
        ]
    )
    if not output.exists() or len(PdfReader(str(output), strict=False).pages) == 0:
        raise RuntimeError(f"LibreOffice did not finish a valid PDF at {output}")
    run(
        [
            "exiftool",
            "-overwrite_original",
            f"-Author={AUTHOR}",
            f"-Title={TITLE}",
            str(output),
        ]
    )
    return output


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for block in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(block)
    return digest.hexdigest()


def build_release(markdown: Path, docx: Path, pdf: Path) -> tuple[Path, Path]:
    attachments = OUTPUT / "attachments"
    attachments.mkdir(parents=True, exist_ok=True)
    archive = attachments / "企业AI系统落地-随书附件-第四稿-2026-07-22.zip"
    archive.unlink(missing_ok=True)

    members = [
        BOOK / "README.md",
        BOOK / "LICENSE-CODE.md",
        BOOK / "appendices" / "A-project-toolkit.md",
        BOOK / "appendices" / "B-glossary.md",
        BOOK / "appendices" / "C-versioned-technical-cards.md",
        BOOK / "appendices" / "D-code-companion.md",
        *sorted((BOOK / "tools").glob("*.md")),
        *sorted((BOOK / "code").glob("*.py")),
        *sorted((BOOK / "code" / "tests").glob("*.py")),
    ]
    inventory = TMP / "随书附件清单.txt"
    inventory_lines = [
        "《企业 AI 系统落地》随书附件",
        "版本：第四稿（2026-07-22）",
        "代码运行：Python 3.10+，仅使用标准库",
        "测试：python3 -m unittest discover -s code/tests -v",
        "",
        "文件 SHA-256：",
    ]
    for member in members:
        inventory_lines.append(f"{sha256(member)}  {member.relative_to(BOOK)}")
    inventory.write_text("\n".join(inventory_lines) + "\n", encoding="utf-8")

    with zipfile.ZipFile(archive, "w", compression=zipfile.ZIP_DEFLATED) as bundle:
        bundle.write(inventory, "随书附件清单.txt")
        for member in members:
            bundle.write(member, member.relative_to(BOOK))

    manifest = OUTPUT / "第四稿-SHA256SUMS.txt"
    released = [markdown, docx, pdf, archive]
    manifest.write_text(
        "".join(f"{sha256(path)}  {path.relative_to(OUTPUT)}\n" for path in released),
        encoding="utf-8",
    )
    return archive, manifest


def main() -> None:
    if TMP.exists():
        shutil.rmtree(TMP)
    markdown = build_markdown()
    docx = build_docx(markdown)
    pdf = build_pdf_from_docx(docx)
    archive, manifest = build_release(markdown, docx, pdf)
    print(markdown)
    print(docx)
    print(pdf)
    print(archive)
    print(manifest)


if __name__ == "__main__":
    main()
